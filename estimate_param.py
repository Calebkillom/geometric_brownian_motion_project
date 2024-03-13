"""
 Reads historical stock price data.
 Calculates log returns.
 Computes the mean and standard deviation of the log returns,
 Using both pandas and NumPy methods.
 Finally, it Saves the results in a word Document
"""
import pandas as pd
import numpy as np
from docx import Document
import os


file_path = "XOM_historical_data_2018-03-09_2024-03-10.csv"

historical_data = pd.read_csv(file_path)

print("Historical data")
print(historical_data.tail())

""" Calculate the natural logarithm of the adjusted close values """
historical_data['Log_Adjusted_Close'] = np.log(historical_data["Adj Close"])

""" Shift the adjusted close values by one row """
historical_data["Prev_Log_Adj_Close"] =\
    historical_data["Log_Adjusted_Close"].shift(1)

"""
Calculate the difference between the natural logarithm
of the adjusted close values of each row and its previous row
this will be the log returns
"""
historical_data["Log Returns"] = historical_data["Log_Adjusted_Close"]\
    - historical_data["Prev_Log_Adj_Close"]

""" Drop the first row which will have NaN due to the shift """
historical_data = historical_data.dropna()

print(historical_data.tail())

"""
Calculating mu which is the mean returns
of the stock prices within the historical data range selected
"""
mu = historical_data["Log Returns"].mean()
print("Mean of log returns:", mu)

mu1 = np.mean(historical_data["Log Returns"])
print("Mean of log returns:", mu1)

"""
Calculating sigma which is the standard deviation
of the return of  stock prices within the historical data range selected
"""
sigma = np.std(historical_data["Log Returns"])
print("The std dev of log returns is:", sigma)

sigma1 = historical_data["Log Returns"].std()
print("The std dev of log returns is:", sigma1)

""" Saving the Results in a Word Document """
""" Creating a new Document Object """
doc = Document()

""" Adding a Heading """
doc.add_heading("Geometric Brownian Motion Q3a) solution", level=1)

""" Adding the mu and Sigma values to the Document """
doc.add_paragraph(f"Mean of log returns: {mu1}")
doc.add_paragraph(f"Standard deviation of log returns: {sigma}")

""" Save the Document """
doc.save("gbm.docx")

""" Check if Document exists """
if os.path.exists("gbm.docx"):
    print("Document Saved Successfully")
